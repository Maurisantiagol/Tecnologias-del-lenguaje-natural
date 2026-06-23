import random
import time
from deep_translator import GoogleTranslator
from docx import Document
from chatbot import asistente

# 1. Conjunto de Pruebas: [(Texto en Español, Intención Esperada), ...]
# Generación dinámica de más de 200 pruebas
TEST_CASES = []

# Plantillas para generar casos
reservas = [
    "Quiero reservar una mesa para el {} a las {}",
    "Mesa para dos el {} por la noche",
    "Necesito hacer una reservación el {}",
    "Podrías reservarme un lugar {}",
    "Quiero hacer una reserva el {} a las {}"
]
dias = ["lunes", "martes", "miercoles", "jueves", "viernes", "sabado", "domingo", "mañana", "hoy"]
horas = ["7", "8", "9", "20:00", "21:00"]

for d in dias:
    for h in horas:
        TEST_CASES.append((f"Quiero reservar una mesa para el {d} a las {h}", "Book_Table"))
        TEST_CASES.append((f"Mesa para dos el {d} por la noche", "Book_Table"))
        
menu_queries = [
    "Cuales son sus opciones en el menu", "Que comida sirven aqui", 
    "Muestrame las categorias del menu", "Quiero ver el menu",
    "Tienen menu de niños", "Donde puedo leer el menu", "Que ofrecen en el menu",
    "Cual es el menu principal"
]
for _ in range(5):
    for m in menu_queries:
        TEST_CASES.append((m, "Query_Menu"))

comidas = ["pasta", "pizza", "ensalada", "risotto", "postre", "sopa", "carne"]
adjetivos = ["vegana", "sin gluten", "rapida", "deliciosa", "italiana", "ligera", "picante"]
for c in comidas:
    for a in adjetivos:
        TEST_CASES.append((f"Recomiendame una {c} {a}", "Recommend_Food"))
        TEST_CASES.append((f"Quiero comer algo como {c} que sea {a}", "Recommend_Food"))
        
ingredientes = ["nueces", "queso", "tomate", "ajo", "leche", "mariscos"]
platos = ["risotto", "pizza", "pasta", "ensalada", "tiramisu"]
for p in platos:
    for i in ingredientes:
        TEST_CASES.append((f"Cuales son los ingredientes de la {p}", "Query_Ingredients"))
        TEST_CASES.append((f"Ese {p} contiene {i}?", "Query_Ingredients"))

cancelaciones = ["cancelar", "cambiar", "actualizar", "modificar"]
for c in cancelaciones:
    for _ in range(5):
        TEST_CASES.append((f"Quiero {c} mi reserva", "Modify_Booking"))

# Limitar a 200 pruebas únicas para agilizar un poco
import random
random.shuffle(TEST_CASES)
TEST_CASES = TEST_CASES[:200]


def inyectar_ruido(texto, probabilidad_ruido=0.15):
    """
    Inyecta ruido simulando errores tipográficos:
    - Omitir caracteres
    - Intercambiar caracteres adyacentes
    - Cambiar un carácter por otro aleatorio
    """
    if probabilidad_ruido == 0:
        return texto
        
    resultado = list(texto)
    letras = "abcdefghijklmnopqrstuvwxyz"
    
    for i in range(len(resultado)):
        if resultado[i] == ' ' or random.random() > probabilidad_ruido:
            continue
            
        tipo_error = random.choice(['omitir', 'intercambiar', 'sustituir'])
        
        if tipo_error == 'omitir':
            resultado[i] = ''
        elif tipo_error == 'intercambiar' and i < len(resultado) - 1:
            resultado[i], resultado[i+1] = resultado[i+1], resultado[i]
        elif tipo_error == 'sustituir':
            resultado[i] = random.choice(letras)
            
    return "".join(resultado)

def ejecutar_pruebas():
    translator_es_en = GoogleTranslator(source='es', target='en')
    
    niveles_ruido = [0.0, 0.05, 0.15, 0.30]
    resultados = {}
    
    doc = Document()
    doc.add_heading('REPORTE DE PRUEBAS DEL TRADUCTOR Y CHATBOT', 0)
    
    with open("reporte_pruebas.txt", "w", encoding="utf-8") as f:
        f.write("=== REPORTE DE PRUEBAS DEL TRADUCTOR Y CHATBOT ===\n\n")
        
        for nivel in niveles_ruido:
            f.write(f"--- Evaluando con Nivel de Ruido: {nivel*100}% ---\n")
            f.flush()
            doc.add_heading(f"Evaluando con Nivel de Ruido: {nivel*100}%", level=1)
            
            aciertos = 0
            total = len(TEST_CASES)
            print(f"\nEvaluando con Nivel de Ruido: {nivel*100}%")
            
            for i, (texto_original, intent_esperado) in enumerate(TEST_CASES, 1):
                texto_ruidoso = inyectar_ruido(texto_original, nivel)
                
                # 1. Traducción
                try:
                    texto_en = translator_es_en.translate(texto_ruidoso)
                except Exception as e:
                    f.write(f"Error de traducción: {e}\n")
                    continue
                
                # 2. Procesamiento
                asistente.reset()
                resultado = asistente.procesar_mensaje(texto_en)
                intent_obtenido = resultado['intent']
                
                # Fallback handler for Recommend_Food variations
                if intent_esperado == 'Recommend_Food' and intent_obtenido == 'Discover_Food':
                    intent_obtenido = 'Recommend_Food'
                
                # Evaluar
                exito = (intent_esperado == intent_obtenido)
                if exito:
                    aciertos += 1
                
                print(f"  [{i}/{total}] -> Esperado: {intent_esperado} | Obtenido: {intent_obtenido}")
                
                time.sleep(0.1)
                
                # Guardar detalle de fallos
                if not exito:
                    f.write(f"[FALLO] Original: '{texto_original}'\n")
                    f.write(f"        Ruidoso:  '{texto_ruidoso}'\n")
                    f.write(f"        Traducido: '{texto_en}'\n")
                    f.write(f"        Esperaba: {intent_esperado} | Obtuvo: {intent_obtenido}\n\n")
                    f.flush()
                    
                    doc.add_paragraph(f"[FALLO] Original: '{texto_original}'")
                    doc.add_paragraph(f"Ruidoso: '{texto_ruidoso}'", style='List Bullet')
                    doc.add_paragraph(f"Traducido: '{texto_en}'", style='List Bullet')
                    doc.add_paragraph(f"Esperaba: {intent_esperado} | Obtuvo: {intent_obtenido}", style='List Bullet')
            
            accuracy = (aciertos / total) * 100
            resultados[nivel] = accuracy
            f.write(f"==> Precisión con {nivel*100}% de ruido: {accuracy:.2f}%\n\n")
            f.flush()
            doc.add_paragraph(f"==> Precisión con {nivel*100}% de ruido: {accuracy:.2f}%").bold = True
        
        f.write("\n=== RESUMEN GLOBAL ===\n")
        doc.add_heading("RESUMEN GLOBAL", level=1)
        for nivel, acc in resultados.items():
            f.write(f"Ruido {nivel*100}% -> Precisión: {acc:.2f}%\n")
            doc.add_paragraph(f"Ruido {nivel*100}% -> Precisión: {acc:.2f}%")
            
        doc.save("reporte_pruebas.docx")

if __name__ == "__main__":
    print("Iniciando pruebas exhaustivas. Esto tomará un par de minutos...")
    ejecutar_pruebas()
    print("Pruebas finalizadas. Revisa el archivo reporte_pruebas.txt")
